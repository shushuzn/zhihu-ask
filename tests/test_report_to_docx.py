"""report_to_docx.py 回归测试：md→docx 转换纯函数 + 端到端（30 项）。

覆盖：
- split_rich：**加粗** 段内标记解析（单段/混合/多段/边界/无标记/空串）
- parse_table_rows：分隔行过滤与单元格拆分（strip/不等宽/空行）
- normalize_img_ext：URL 扩展名推导（查询串/白名单/未知回退/大写归一）
- convert_md_to_docx 端到端（python-docx 可用时）：标题层级、段落文本、
  加粗 run、bullet/有序/引用、表格单元格、行内公式 $...$ → 默认 omml 模式
  （Word 原生公式）/ text 模式（Unicode 文本）/ image 模式（PNG）/ svg 模式
  （SVG 矢量+PNG 回退）——锁定「只转格式一字不改」契约

运行：python tests/test_report_to_docx.py
"""
import os
import sys
import tempfile
import shutil

import testutil

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(ROOT, "tools"))
import report_to_docx as r2d

PASS = 0
FAIL = 0


def expect(label, got, must_be):
    global PASS, FAIL
    if got == must_be:
        PASS += 1
    else:
        FAIL += 1
        print(f"  FAIL {label}: got {got!r}, expected {must_be!r}")


# ---- split_rich：加粗解析 ----
expect("rich+ 纯文本", r2d.split_rich("普通文本"), [("普通文本", False)])
expect("rich+ 整段加粗", r2d.split_rich("**加粗词**"), [("加粗词", True)])
expect("rich+ 混合", r2d.split_rich("前**中**后"),
       [("前", False), ("中", True), ("后", False)])
expect("rich+ 多段加粗", r2d.split_rich("**a** x **b**"),
       [("a", True), (" x ", False), ("b", True)])
expect("rich+ 开头加粗", r2d.split_rich("**首**尾"),
       [("首", True), ("尾", False)])
expect("rich+ 空串", r2d.split_rich(""), [])
expect("rich- 单星号不解析", r2d.split_rich("*星号*"), [("*星号*", False)])
expect("rich- 裸双星号不解析", r2d.split_rich("a**b"), [("a**b", False)])
expect("rich- 三连星不解析", r2d.split_rich("a***b"), [("a***b", False)])

# ---- parse_table_rows：表格单元格 ----
rows = ["| 列A | 列B |", "|---|---|", "| 甲 | 乙 |"]
expect("table+ 分隔行过滤", r2d.parse_table_rows(rows), [["列A", "列B"], ["甲", "乙"]])
rows = ["| a | b | c |", "|--|--|--|", "| 1 | 2 | 3 |"]
expect("table+ 三列", r2d.parse_table_rows(rows), [["a", "b", "c"], ["1", "2", "3"]])
rows = ["| 左 | 右 |", "|:--|--:|", "| x | y |"]
expect("table+ 对齐分隔行也过滤", r2d.parse_table_rows(rows), [["左", "右"], ["x", "y"]])
rows = ["| 不等宽 |", "|---|", "| 1 | 2 |"]
expect("table+ 不等宽保留原样", r2d.parse_table_rows(rows), [["不等宽"], ["1", "2"]])
expect("table+ 空列表", r2d.parse_table_rows([]), [])
rows = ["|---|---|"]
expect("table+ 仅分隔行", r2d.parse_table_rows(rows), [])
# 表格内 LaTeX 范数 \\| 不得被误判为列分隔符
rows = ["| 风险 | 说明 |", "|---|---|",
        '| 数值猜测未证实 | $\\|T_{\\mathrm{even}}\\|\\approx 0.6368$ 与 $\\alpha\\beta=\\pi$ |']
expect("table+ 公式内竖线保护",
       r2d.parse_table_rows(rows),
       [["风险", "说明"],
        ["数值猜测未证实", "$\\|T_{\\mathrm{even}}\\|\\approx 0.6368$ 与 $\\alpha\\beta=\\pi$"]])

# ---- normalize_img_ext：扩展名 ----
expect("ext+ png 带查询串", r2d.normalize_img_ext("https://x/a.png?raw=1&v=2"), ".png")
expect("ext+ jpeg", r2d.normalize_img_ext("https://x/a.jpeg"), ".jpeg")
expect("ext+ gif", r2d.normalize_img_ext("https://x/a.gif"), ".gif")
expect("ext+ 大写归一", r2d.normalize_img_ext("https://x/a.PNG"), ".png")
expect("ext- webp 回退 png", r2d.normalize_img_ext("https://x/a.webp"), ".png")
expect("ext- 无扩展名回退", r2d.normalize_img_ext("https://x/image"), ".png")
expect("ext- 未知扩展回退", r2d.normalize_img_ext("https://x/a.tiff"), ".png")

# ---- convert_md_to_docx 端到端（锁定一字不改契约） ----
try:
    import docx as _docx_mod
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

if _HAS_DOCX:
    from docx import Document
    tmp = testutil.mktestdir()
    try:
        md = (
            "# 标题一\n\n"
            "正文段落包含**加粗词**和普通词，以及行内公式 $A_2A_4 - A_3^2$ 混排。\n\n"
            # mathml2omml 对 \bar 漏闭合 groupChrPr，
            # latex_to_omml 已把 \bar→\overline 规避，含 \bar 公式必须转成 OMML 而非回退文本
            "共轭内积公式 $\\langle f,g\\rangle=\\int f\\bar{g}$ 应转 OMML。\n\n"
            "- 列表项一\n"
            "- 列表项二\n\n"
            "1. 有序项\n\n"
            "> 引用内容\n\n"
            "| 列A | 列B |\n"
            "|---|---|\n"
            "| 甲 | 乙 |\n\n"
            "## 小节二\n\n"
            "结尾。\n"
        )
        md_path = os.path.join(tmp, "report.md")
        out_path = os.path.join(tmp, "report.docx")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md)
        r2d.convert_md_to_docx(md_path, out_path)
        expect("docx+ 文件生成", os.path.exists(out_path), True)

        d = Document(out_path)
        texts = [p.text for p in d.paragraphs]

        expect("docx+ 标题一为 Heading1",
               any(p.style.name == "Heading 1" and p.text == "标题一" for p in d.paragraphs), True)
        expect("docx+ 小节二为 Heading2",
               any(p.style.name == "Heading 2" and p.text == "小节二" for p in d.paragraphs), True)
        expect("docx+ 加粗标记剥离",
               any("正文段落包含加粗词和普通词，以及行内公式" in t for t in texts)
               and any("混排。" in t for t in texts), True)
        bold_run = None
        for p in d.paragraphs:
            for r in p.runs:
                if r.text == "加粗词" and r.font.bold:
                    bold_run = r
        expect("docx+ 加粗 run 生效", bold_run is not None, True)
        expect("docx+ bullet 项",
               any(p.style.name == "List Bullet" and p.text == "列表项一" for p in d.paragraphs), True)
        expect("docx+ 有序项",
               any(p.style.name == "List Number" and p.text == "有序项" for p in d.paragraphs), True)
        expect("docx+ 引用内容", "引用内容" in texts, True)
        expect("docx+ 表格 2 行 2 列", len(d.tables) == 1 and len(d.tables[0].rows) == 2
               and len(d.tables[0].columns) == 2, True)
        expect("docx+ 表格单元格", d.tables[0].cell(0, 0).text == "列A"
               and d.tables[0].cell(1, 1).text == "乙", True)
        # 一字不改：所有文字拼接后不含 markdown 语法残留
        joined = "".join(texts)
        expect("docx+ 无 md 语法残留",
               "**" not in joined and "# 标题" not in joined and "|" not in joined, True)
        # 默认 omml 模式：LaTeX 公式 → Word 原生公式（本地显示排版公式且可编辑）
        import zipfile as _zipfile
        with _zipfile.ZipFile(out_path) as zf:
            names = zf.namelist()
            xml = zf.read("word/document.xml").decode("utf-8")
        expect("docx+ 公式转 OMML", "<m:oMath>" in xml, True)
        expect("docx+ 无 LaTeX 源码残留", "$A_2A_4" not in xml, True)
        # 含 \bar 的公式必须转 OMML，不得回退为 $...$ 纯文本
        expect("docx+ bar 公式转 OMML", r"$\langle f,g\rangle=\int f\bar{g}$" not in xml, True)
        # text 模式：LaTeX → Unicode 数学文本
        out_text = os.path.join(tmp, "report_text.docx")
        r2d.convert_md_to_docx(md_path, out_text, formula_mode="text")
        with _zipfile.ZipFile(out_text) as zf:
            xml_text = zf.read("word/document.xml").decode("utf-8")
        expect("docx+ text 模式公式转 Unicode", "A₂A₄" in xml_text, True)
        # image 模式：LaTeX → PNG 图片嵌入
        out_img = os.path.join(tmp, "report_img.docx")
        r2d.convert_md_to_docx(md_path, out_img, formula_mode="image")
        with _zipfile.ZipFile(out_img) as zf:
            names_img = zf.namelist()
        expect("docx+ image 模式公式转图片", any(n.startswith("word/media/") for n in names_img), True)
        # svg 模式：LaTeX → SVG 矢量 + PNG 回退双轨
        out_svg = os.path.join(tmp, "report_svg.docx")
        r2d.convert_md_to_docx(md_path, out_svg, formula_mode="svg")
        with _zipfile.ZipFile(out_svg) as zf:
            names_svg = zf.namelist()
            xml_svg = zf.read("word/document.xml").decode("utf-8")
        expect("docx+ svg 模式含 SVG 公式图", any(n.endswith(".svg") for n in names_svg), True)
        expect("docx+ svg 模式含 svgBlip 引用", "svgBlip" in xml_svg, True)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)
else:
    print("  [跳过] 当前解释器无 python-docx，端到端用例未执行")


print(f"TOTAL: PASS={PASS} FAIL={FAIL}")
