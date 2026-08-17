
"""
研究报告 → docx（Word）转换工具（zhihu-ask 项目专用，）

把 report.md（URL 引用版，图片走公网 https URL）转换为 Word 文档：
  - 标题（#/##/###）→ Word 标题样式（Heading 1/2/3），中文字体微软雅黑/宋体
  - 段落/列表（- 、1.）→ 正文/列表样式
  - 表格（| a | b |）→ Word 表格
  - 图片 ![alt](https://...) → 下载公网图片并嵌入文档（失败则保留 URL 文本）
  - 图注（图 N｜说明）→ 图片下方居中灰色小字
  - 加粗 **text** → 加粗 run
  - 行内公式 $...$ → 默认 Word 原生公式（OMML，LaTeX→MathML→OMML，本地显示排版
    公式且可编辑）；`--formula-mode svg`（SVG 矢量+PNG 回退）/ `text`（Unicode 文本，
    上传最稳）/ `image`（PNG）可切换
输出：research/<slug>/report.docx（与 report.md 同级，单一文件约定）

依赖：python-docx（text 模式）；PIL、matplotlib（image 模式）；latex2mathml、mathml2omml（omml 模式）
（隔离 venv，默认仓库根 `venv/`，可用环境变量 `ZHIHU_ASK_VENV_PY` 覆盖）；
主解释器无依赖时自动切换 venv 重跑（同 report_images.py 的 ensure_venv 模式）。

用法：
    python tools/report_to_docx.py --slug <slug>              # 生成 research/<slug>/report.docx
    python tools/report_to_docx.py --slug <slug> --out x.docx  # 自定义输出文件名（默认 report.docx）

说明：
- 图片 URL 需公网可访问（report_images.py --url-base 部署的 CloudStudio 静态托管）；
  下载失败时图片位置保留 alt（URL）文本，不中断转换。
- 转换只做格式映射，不增删任何文字（同 report_to_flomo 的一字不改原则）。
"""

import argparse
import os
import re
import subprocess
import sys
import tempfile
import urllib.request

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# 外网出口检测（沙箱 Bash 通常无 egress，图片下载等外网动作需先确认）
try:
    from net_check import has_egress as _egress_ok
except Exception:  # 兜底：无法导入时按"有出口"处理，保留原有异常路径
    def _egress_ok():
        return True
# 隔离 venv 的 Python 解释器；可用环境变量 ZHIHU_ASK_VENV_PY 覆盖，
# 未设置时回退到仓库根 venv/Scripts/python.exe（Windows）或 venv/bin/python（类 Unix）。
if os.name == "nt":
    _DEFAULT_VENV_PY = os.path.join(ROOT, "venv", "Scripts", "python.exe")
else:
    _DEFAULT_VENV_PY = os.path.join(ROOT, "venv", "bin", "python")
VENV_PY = os.environ.get("ZHIHU_ASK_VENV_PY", _DEFAULT_VENV_PY)

EAST_ASIA_FONT = "微软雅黑"
ASCII_FONT = "Microsoft YaHei"

def ensure_docx():
    """当前解释器缺依赖时，自动创建隔离 venv 并安装后重跑自身。"""
    try:
        import docx  # noqa: F401
        import PIL  # noqa: F401
        import matplotlib  # noqa: F401
        return
    except ImportError:
        pass
    # 复用已有 venv
    if not os.path.exists(VENV_PY):
        print("主解释器缺依赖，正在创建隔离 venv 并安装…")
        import venv
        # venv 根目录是 VENV_PY 的上两级（venv/Scripts/python.exe），
        # 不能传 Scripts 层，否则把 venv 嵌套建到 venv/Scripts/ 下导致 python.exe 找不到
        venv.create(os.path.dirname(os.path.dirname(VENV_PY)), with_pip=True)
        pip = [VENV_PY, "-m", "pip", "install", "-q",
               "python-docx", "Pillow", "matplotlib", "latex2mathml", "mathml2omml"]
        try:
            subprocess.run(pip, check=True)
        except subprocess.CalledProcessError as e:
            print(f"ERROR: venv 内安装依赖失败：{e}")
            sys.exit(1)
    print("切换到隔离 venv 重跑…")
    os.execv(VENV_PY, [VENV_PY] + sys.argv)

def normalize_img_ext(url):
    """纯函数：从图片 URL 推导扩展名（白名单外的回退 .png）。"""
    ext = os.path.splitext(url.split("?")[0])[1].lower() or ".png"
    if ext not in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
        ext = ".png"
    return ext

def download_image(url, timeout=60, base_dir=None):
    """下载公网图片到临时文件，返回 (临时路径, 扩展名)；失败返回 None。
    支持本地路径（含相对路径，需传 base_dir=报告目录才能解析相对引用）。"""
    ext = normalize_img_ext(url)
    try:
        is_remote = url.startswith(("http://", "https://", "file://", "data:"))
        resolved = url
        if not is_remote and base_dir:
            resolved = os.path.join(base_dir, url)
        if not is_remote and os.path.isfile(resolved):
            data = open(resolved, "rb").read()
        else:
            # urllib 经 HTTPS_PROXY 通常可联网；若探测失败（无 egress）则跳过并提示
            if not _egress_ok():
                print("    [图片下载跳过] 当前环境 urllib 无外网出口，图片保留 URL 文本。")
                return None, None
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                data = resp.read()
        fd, tmp = tempfile.mkstemp(suffix=ext)
        with os.fdopen(fd, "wb") as f:
            f.write(data)
        return tmp, ext
    except Exception as e:
        print(f"    [图片下载失败] {e}")
        return None, None

def set_run_font(run, size=None, bold=None, color=None):
    """设置 run 的中英文字体。"""
    from docx.oxml.ns import qn
    run.font.name = ASCII_FONT
    r = run._element.rPr
    if r is not None:
        rFonts = r.find(qn("w:rFonts"))
        if rFonts is not None:
            rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color

def _normalize_frac_args(latex):
    """把 \frac 的单字符简写参数补成花括号形式（mathtext 要求）。

    覆盖三种形态：\frac12 → \frac{1}{2}；\frac1{12} → \frac{1}{12}；
    \frac{12}4 → \frac{12}{4}。逐 token 扫描 \frac 后的两个参数。
    """
    out = []
    i = 0
    n = len(latex)
    while i < n:
        if latex.startswith("\\frac", i):
            out.append("\\frac")
            i += 5
            # 解析两个参数
            for _ in range(2):
                while i < n and latex[i] == " ":
                    i += 1
                if i < n and latex[i] == "{":
                    # 花括号参数：找配对 }
                    depth = 0
                    j = i
                    while j < n:
                        if latex[j] == "{":
                            depth += 1
                        elif latex[j] == "}":
                            depth -= 1
                            if depth == 0:
                                j += 1
                                break
                        j += 1
                    out.append(latex[i:j])
                    i = j
                else:
                    # 单字符参数 → 补花括号
                    out.append("{" + latex[i] + "}")
                    i += 1
        else:
            out.append(latex[i])
            i += 1
    return "".join(out)


def latex_to_text(latex):
    """LaTeX 公式 → Unicode 数学文本（docx 默认公式模式）。

    转换：分数 → (a)/(b)；\\tfrac/\\dfrac → \\frac 并补花括号；\\beta → β 等常见命令；
    数字/逗号下标 → Unicode 下标（A_2 → A₂、H_{2,2} → H₂,₂）、数字上标 → Unicode 上标
    （x^2 → x²）；非数字上下标 → ^(...)/_(...)。
    任何平台把结果当普通文本显示——不依赖 OMML/图片，排版与正文一致。
    """
    s = latex
    # 1) 分数统一：\tfrac/\dfrac → \frac，简写参数补花括号
    s = s.replace("\\tfrac", "\\frac").replace("\\dfrac", "\\frac")
    s = _normalize_frac_args(s)
    # 2) 分数 → (a)/(b)，由内向外循环（支持嵌套）
    for _ in range(20):
        m = re.search(r"\\frac\{([^{}]*)\}\{([^{}]*)\}", s)
        if not m:
            break
        s = s[:m.start()] + f"({m.group(1)})/({m.group(2)})" + s[m.end():]
    # 3) 命令 → Unicode/文本（长命令先替换，防 \cdot 截胡 \cdots）
    s = s.replace("\\cdots", "…").replace("\\ldots", "…").replace("\\dots", "…")
    s = s.replace("\\left(", "(").replace("\\right)", ")").replace("\\left[", "[")
    s = s.replace("\\right]", "]").replace("\\,", " ").replace("\\;", " ").replace("\\!", "")
    for cmd, sym in [
        ("\\beta", "β"), ("\\alpha", "α"), ("\\gamma", "γ"), ("\\pi", "π"),
        ("\\theta", "θ"), ("\\phi", "φ"), ("\\mu", "μ"), ("\\sigma", "σ"),
        ("\\lambda", "λ"), ("\\delta", "δ"), ("\\omega", "ω"),
        ("\\cos", "cos"), ("\\sin", "sin"), ("\\tan", "tan"), ("\\sec", "sec"),
        ("\\csc", "csc"), ("\\cot", "cot"), ("\\log", "log"), ("\\ln", "ln"),
        ("\\exp", "exp"), ("\\leq", "≤"), ("\\geq", "≥"), ("\\neq", "≠"),
        ("\\approx", "≈"), ("\\cdot", "·"), ("\\times", "×"), ("\\pm", "±"),
        ("\\infty", "∞"), ("\\in", "∈"), ("\\sqrt", "√"),
    ]:
        s = s.replace(cmd, sym)
    # 4) 下标/上标 → Unicode（纯数字或数字+逗号）
    sub_map = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")
    sup_map = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")
    s = re.sub(r"_\{([0-9,]+)\}", lambda m: m.group(1).translate(sub_map), s)
    s = re.sub(r"\^\{([0-9]+)\}", lambda m: m.group(1).translate(sup_map), s)
    s = re.sub(r"_\{([^}]*)\}", lambda m: "_(" + m.group(1) + ")", s)  # 非纯数字下标兜底
    s = re.sub(r"\^\{([^}]*)\}", lambda m: "^(" + m.group(1) + ")", s)  # 非纯数字上标兜底
    s = re.sub(r"_([0-9])", lambda m: m.group(1).translate(sub_map), s)
    s = re.sub(r"\^([0-9])", lambda m: m.group(1).translate(sup_map), s)
    # 5) 清理残留花括号与未知命令
    s = s.replace("{", "").replace("}", "")
    s = re.sub(r"\\([a-zA-Z]+)", "", s)
    # 压缩多余空格（连字符两侧保留单空格）
    s = re.sub(r"  +", " ", s)
    return s.strip()


def latex_to_png(latex, out_path):
    """LaTeX 公式 → 透明背景 PNG（matplotlib mathtext，无需系统 LaTeX）。

    归一化：\\tfrac/\\dfrac → \\frac；\\frac 单字符参数简写补花括号（mathtext 要求）。
    返回 PNG 路径；渲染失败返回 None（调用方回退纯文本源码）。
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return None
    norm = latex.replace("\\tfrac", "\\frac").replace("\\dfrac", "\\frac")
    norm = _normalize_frac_args(norm)
    try:
        fig = plt.figure(figsize=(6, 1))
        fig.text(0.5, 0.5, "$" + norm + "$", ha="center", va="center", fontsize=14)
        fig.savefig(out_path, dpi=300, bbox_inches="tight", transparent=True)
        plt.close(fig)
        return out_path if os.path.getsize(out_path) > 0 else None
    except Exception:
        try:
            plt.close("all")
        except Exception:
            pass
        return None


def latex_to_svg_png(latex, svg_path, png_path):
    """LaTeX 公式 → 同时渲染 SVG + PNG（matplotlib mathtext）。

    同一公式渲染两份：SVG（矢量，Word 2016+ 显示，任意缩放不模糊）+
    PNG（位图回退，旧版 Word/部分平台用）。返回 (svg_ok, png_ok)。
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return (False, False)
    norm = latex.replace("\\tfrac", "\\frac").replace("\\dfrac", "\\frac")
    norm = _normalize_frac_args(norm)
    svg_ok = png_ok = False
    try:
        fig = plt.figure(figsize=(6, 1))
        fig.text(0.5, 0.5, "$" + norm + "$", ha="center", va="center", fontsize=14)
        fig.savefig(svg_path, format="svg", transparent=True)
        svg_ok = os.path.getsize(svg_path) > 0
        fig.savefig(png_path, format="png", dpi=300, transparent=True)
        png_ok = os.path.getsize(png_path) > 0
        plt.close(fig)
    except Exception:
        try:
            plt.close("all")
        except Exception:
            pass
    return (svg_ok, png_ok)


def add_svg_formula(paragraph, latex, size):
    """LaTeX 公式 → SVG 矢量图 + PNG 回退双轨嵌入段落。

    OOXML 规范：a:blip 主引用 PNG（兼容旧版/平台），extLst 内 svgBlip 引用
    SVG（Word 2016+ 显示矢量）。svgBlip URI 为微软 SVG 扩展固定值。
    返回 True 成功 / False 失败（调用方回退）。
    """
    svg_path = _next_formula_png().replace(".png", ".svg")
    png_path = svg_path.replace(".svg", ".png")
    svg_ok, png_ok = latex_to_svg_png(latex, svg_path, png_path)
    if not (svg_ok and png_ok):
        return False
    try:
        from docx.shared import Inches as _Inches
        from PIL import Image as _PILImage
        with _PILImage.open(png_path) as im:
            w_px, h_px = im.size
        base_pt = size.pt if size else 10.5
        target_h_in = base_pt * 1.3 / 72.0
        ratio = h_px / w_px if w_px else 1.0
        width_in = target_h_in / ratio if ratio > 0 else 1.0

        run = paragraph.add_run()
        run.add_picture(png_path, width=_Inches(min(width_in, 5.5)))
        r = run._r

        # SVG 部件 + 关系
        pkg = paragraph.part.package
        partname = pkg.next_partname("/word/media/image%s.svg")
        from docx.opc.part import Part
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        svg_part = Part(partname, "image/svg+xml", open(svg_path, "rb").read(), pkg)
        r_id = paragraph.part.relate_to(svg_part, RT.IMAGE)

        # 注入 svgBlip
        from docx.oxml import parse_xml
        blip = r.xpath(".//a:blip")[0]
        ext_lst = parse_xml(
            '<a:extLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
            '<a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
            '<asvg:svgBlip xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
            f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:embed="{r_id}"/>'
            "</a:ext></a:extLst>"
        )
        blip.append(ext_lst)
        return True
    except Exception:
        return False


# 公式图片临时目录（convert 结束时清理）
_FORMULA_TMP = os.path.join(tempfile.gettempdir(), "zhihu_formula_img")
_FORMULA_IMG_NO = [0]


def _next_formula_png():
    os.makedirs(_FORMULA_TMP, exist_ok=True)
    _FORMULA_IMG_NO[0] += 1
    return os.path.join(_FORMULA_TMP, f"f{_FORMULA_IMG_NO[0]}.png")


def latex_to_omml(latex):
    """LaTeX 公式 → OMML（Word 原生公式 XML）。

    转换链：latex2mathml（LaTeX→MathML）+ mathml2omml（MathML→OMML）。
    失败时返回 None（调用方回退为纯文本展示 LaTeX 源码）。

    已知缺陷修复（fourier 报告踩坑）：mathml2omml 对 \\bar 生成的
    overline 结构（<m:groupChr><m:groupChrPr>...）漏写 </m:groupChrPr> 闭合，
    parse_xml 报 XMLSyntaxError 导致含 \\bar 的公式全部回退纯文本。将 \\bar 替换为
    \\overline（同义、转换路径正常）从源头规避；再对输出做 groupChrPr 配对防御兜底。
    """
    try:
        from latex2mathml.converter import convert as l2m
        from mathml2omml import convert as m2o
        # mathml2omml 缺陷规避：\bar → \overline（仅替换裸 \bar，不碰 \bar{x} 等已带花括号的）
        latex = re.sub(r"\\bar(?![\w])", r"\\overline", latex)
        mathml = l2m(latex)
        omml = m2o(mathml)
        if omml.count("<m:groupChrPr>") > omml.count("</m:groupChrPr>"):
            omml = _repair_groupchr(omml)
        return omml
    except Exception:
        return None


def _repair_groupchr(omml):
    """兜底修复 mathml2omml 的 groupChrPr 未闭合缺陷（纯函数，供测试）。

    mathml2omml 对 \\bar 生成 <m:groupChr><m:groupChrPr><m:chr .../><m:pos .../>
    </m:groupChr> 时漏写 </m:groupChrPr>（错误地用 </m:groupChr> 收尾 groupChrPr）。
    本函数把该错配纠正：将紧跟 <m:groupChrPr> 属性之后出现的第一个 </m:groupChr>
    前补上 </m:groupChrPr>。\\bar→\\overline 已在 latex_to_omml 源头规避，此为双保险。
    """
    # 只处理含 groupChrPr 且缺闭合的情况
    opens = omml.count("<m:groupChrPr>")
    closes = omml.count("</m:groupChrPr>")
    if opens <= closes:
        return omml
    out = omml
    for _ in range(opens - closes):
        # 找到最后一个未配对的 <m:groupChrPr>，在其后最近的 </m:groupChr> 前补闭合
        idx = out.rfind("<m:groupChrPr>")
        if idx == -1:
            break
        close_idx = out.find("</m:groupChr>", idx + len("<m:groupChrPr>"))
        if close_idx == -1:
            break
        out = out[:close_idx] + "</m:groupChrPr>" + out[close_idx:]
    return out


def split_rich(text):
    """纯函数：解析 **加粗** 段内标记，返回 [(片段, 是否加粗)] 列表。

    规则与旧 add_rich_text 一致：`**非星号内容**` 判为加粗段，其余为普通段；
    空片段跳过；无成对标记时整段为普通段。
    """
    out = []
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            out.append((part[2:-2], True))
        else:
            out.append((part, False))
    return out


def add_rich_text(paragraph, text, base_bold=False, size=None, formula_mode="omml"):
    """解析 **加粗** 段内标记并添加 runs；`$...$` 内联公式按 formula_mode 渲染。

    formula_mode：
      - "omml"（默认）：LaTeX → Word 原生公式（OMML）——本地 Word/WPS 显示为
        真正的排版公式（分式/上下标/根号）且可编辑；部分平台上传时可能剥除
        公式元素导致变空格（本项目默认采用此模式）。
      - "svg"：LaTeX → SVG 矢量图 + PNG 回退双轨嵌入——Word 2016+ 显示矢量
        （任意缩放不模糊），旧版/部分平台回退 PNG；行内排版与 PNG 同类但清晰度更优。
      - "text"：LaTeX → Unicode 数学文本——任何平台（公众号/在线文档/云盘）
        上传 docx 公式都按普通文本显示，排版与正文一致、可复制；代价是分式等
        以 (a)/(b) 文本形式呈现，非排版样式。
      - "image"：LaTeX → PNG 图片嵌入——按公式自然宽高比缩放（目标高度约
        正文字号的 1.3 倍），视觉接近 LaTeX 排版；但用户实测上传后行内图片
        排版会乱，仅作备选。
    所有模式转换失败均回退为纯文本展示 LaTeX 源码（不中断、不丢内容）。
    """
    parts = re.split(r"(\$\$[^$\n]+\$\$|\$[^$\n]+\$)", text)
    for part in parts:
        if not part:
            continue
        formula = None
        if part.startswith("$$") and part.endswith("$$") and len(part) > 4:
            formula = part[2:-2]
        elif part.startswith("$") and part.endswith("$") and len(part) > 2:
            formula = part[1:-1]
        if formula is not None:
            latex = formula
            if formula_mode == "omml":
                omml = latex_to_omml(latex)
                if omml:
                    try:
                        from docx.oxml import parse_xml
                        # mathml2omml 输出缺 m 命名空间声明，解析前补上
                        omml = omml.replace(
                            "<m:oMath>",
                            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">',
                            1,
                        )
                        omath = parse_xml(omml)
                        paragraph._p.append(omath)
                        continue
                    except Exception:
                        pass  # 插入失败回退纯文本
            elif formula_mode == "image":
                png = latex_to_png(latex, _next_formula_png())
                if png:
                    try:
                        from docx.shared import Inches as _Inches
                        from PIL import Image as _PILImage
                        with _PILImage.open(png) as im:
                            w_px, h_px = im.size
                        # 目标高度：与字号匹配（正文字号 pt → 高度约 1.3 倍字号），
                        # 按宽高比反推宽度，避免短公式拉伸/大公式撑行
                        base_pt = size.pt if size else 10.5
                        target_h_in = base_pt * 1.3 / 72.0
                        ratio = h_px / w_px if w_px else 1.0
                        width_in = target_h_in / ratio if ratio > 0 else 1.0
                        run = paragraph.add_run()
                        run.add_picture(png, width=_Inches(min(width_in, 5.5)))
                        continue
                    except Exception:
                        pass  # 插入失败回退纯文本
            elif formula_mode == "svg":
                if add_svg_formula(paragraph, latex, size):
                    continue
            else:  # text
                txt = latex_to_text(latex)
                if txt:
                    run = paragraph.add_run(txt)
                    set_run_font(run, size=size, bold=False)
                    continue
        for sub, is_bold in split_rich(part):
            run = paragraph.add_run(sub)
            set_run_font(run, size=size, bold=is_bold or base_bold)

def parse_table_rows(rows):
    """纯函数：过滤分隔行并拆分为单元格矩阵（strip 空格）。

    分隔行形如 `| --- | :--: |`；返回数据行的单元格列表。
    保护 $...$ 内联公式中的 `|`（如 LaTeX 范数 \\|），避免被误判为列分隔符
    （kato 报告踩坑：表格内 $\\|T_{\\mathrm{even}}\\|$ 拆坏单元格）。
    """
    data_rows = [r for r in rows if not re.match(r"^\|[\s:\-|]+\|$", r)]
    cells = []
    for r in data_rows:
        # 先把 $...$ 内的 | 换占位符，按 | 拆列，再还原
        protected = re.sub(r"(\$[^$\n]+\$)", lambda m: m.group(1).replace("|", "\x00"), r)
        cols = [c.strip() for c in protected.strip("|").split("|")]
        cells.append([c.replace("\x00", "|") for c in cols])
    return cells

def convert_md_to_docx(md_path, out_path, formula_mode="omml"):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = ASCII_FONT
    style.font.size = Pt(10.5)
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)

    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    # 纵深防御：AI 概念图（ai_*.png）禁止进 report.md 正文，
    # 若发现引用给出警告（quality_check check_cover_ban 已在门禁硬拦，此处仅提示）。
    import re as _re
    for _li, _ln in enumerate(lines, 1):
        if _re.search(r"!\[[^\]]*\]\([^)]*ai[^)]*\.png\)", _ln):
            print(f"[警告] 行{_li}: report.md 含 AI 概念图引用（{_ln.strip()[:50]}）——"
                  f"封面仅作独立文件 ai_cover.png，禁止进正文，请删除该行。")

    img_fail_alt = []

    i = 0
    while i < len(lines):
        s = lines[i].rstrip()

        if not s.strip():
            i += 1
            continue

        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", s.strip())
        if m:
            alt, url = m.group(1), m.group(2)
            tmp, ext = download_image(url, base_dir=os.path.dirname(md_path))
            if tmp:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(tmp, width=Inches(5.8))
                    os.unlink(tmp)
                except Exception as e:
                    print(f"    [图片嵌入失败] {e}")
                    img_fail_alt.append(alt)
            else:
                img_fail_alt.append(alt)
            i += 1
            continue

        m = re.match(r"^图\s*\d+\s*｜(.+)$", s.strip())
        if m:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_text(p, f"图 {m.group(1).strip()}" if not s.strip().startswith("图 ") else s.strip(), size=Pt(9), formula_mode=formula_mode)
            for run in p.runs:
                set_run_font(run, size=Pt(9), color=RGBColor(0x66, 0x66, 0x66))
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.+)$", s)
        if m:
            level = min(len(m.group(1)), 3)
            heading = doc.add_heading("", level=level)
            run = heading.add_run(m.group(2).strip())
            set_run_font(run, size=[Pt(16), Pt(14), Pt(12)][level - 1], bold=True)

            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
            i += 1
            continue

        if s.strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1

            cells_list = parse_table_rows(rows)
            if cells_list:
                ncol = max(len(r) for r in cells_list)
                table = doc.add_table(rows=len(cells_list), cols=ncol)
                table.style = "Light Grid Accent 1"
                for ri, row in enumerate(cells_list):
                    for ci in range(ncol):
                        cell_text = row[ci] if ci < len(row) else ""
                        cell = table.cell(ri, ci)
                        cell.text = ""
                        add_rich_text(cell.paragraphs[0], cell_text, size=Pt(10), formula_mode=formula_mode)
            continue

        m = re.match(r"^\s*[-*]\s+(.+)$", s)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, m.group(1), size=Pt(10.5), formula_mode=formula_mode)
            i += 1
            continue

        m = re.match(r"^\s*(\d+)[.、]\s+(.+)$", s)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, m.group(2), size=Pt(10.5), formula_mode=formula_mode)
            i += 1
            continue

        if s.startswith(">"):
            p = doc.add_paragraph()
            add_rich_text(p, s.lstrip("> ").strip(), size=Pt(10.5), formula_mode=formula_mode)
            i += 1
            continue

        p = doc.add_paragraph()
        add_rich_text(p, s, size=Pt(10.5), formula_mode=formula_mode)
        i += 1

    doc.save(out_path)
    print(f"已生成 Word 文档: {out_path}（{os.path.getsize(out_path) // 1024} KB）")
    if img_fail_alt:
        print(f"  [注意] {len(img_fail_alt)} 张图片下载失败，文档中以 URL 文本保留:")
        for alt in img_fail_alt:
            print(f"    - {alt}")

def main():
    ap = argparse.ArgumentParser(description="研究报告 → docx 转换工具")
    ap.add_argument("--slug", help="研究报告 slug（research/<slug>/report.md）")
    ap.add_argument("--file", help="直接指定 report.md 路径")
    ap.add_argument("--out", help="输出文件名（默认 report.docx，写 research/<slug>/）")
    ap.add_argument("--formula-mode", choices=["text", "image", "svg", "omml"], default="omml",
                    help="公式渲染模式：omml=Word 原生公式（默认，本地显示排版公式且可编辑）；"
                         "svg=SVG 矢量图+PNG 回退双轨（Word 2016+ 显示矢量、清晰度优）；"
                         "text=Unicode 文本（上传任何平台排版最稳）；"
                         "image=PNG 图片嵌入（视觉接近 LaTeX 排版，但上传后行内图片排版易乱）")
    args = ap.parse_args()

    ensure_docx()

    if args.file:
        md_path = args.file
        slug = os.path.basename(os.path.dirname(args.file))
    elif args.slug:
        slug = args.slug
        md_path = os.path.join(ROOT, "research", slug, "report.md")
    else:
        print("ERROR: 需 --slug 或 --file")
        sys.exit(1)
    if not os.path.isfile(md_path):
        print(f"ERROR: 未找到报告 {md_path}")
        sys.exit(1)

    out_name = args.out or "report.docx"
    if not os.path.isabs(out_name):
        out_path = os.path.join(ROOT, "research", slug, out_name)
    else:
        out_path = out_name

    convert_md_to_docx(md_path, out_path, formula_mode=args.formula_mode)
    # 清理公式渲染临时图片
    if os.path.isdir(_FORMULA_TMP):
        try:
            import shutil
            shutil.rmtree(_FORMULA_TMP, ignore_errors=True)
        except Exception:
            pass

if __name__ == "__main__":
    main()
